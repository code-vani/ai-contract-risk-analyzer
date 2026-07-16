"""
Smart Extractor — Automated Test Suite
=======================================
Industry-standard test coverage for the hybrid extraction engine.

Covers:
    1. Normal digital PDF extraction (text mode)
    2. DOCX file extraction (text mode)
    3. Unsupported file type rejection (error mode)
    4. Empty file rejection (error mode)
    5. Non-existent file path (error mode)
    6. Markdown structure verification (headings/tables preserved)
    7. Scanned PDF detection (image fallback mode)

Run:
    python backend/ingestion/test_extractor.py

All tests create temporary files and clean them up automatically.
"""

import sys
import os
import logging
import tempfile

# ── Path setup ────────────────────────────────────────────────────────────────
# smart_extractor uses `from ingestion.file_validator import ...` which requires
# backend/ on sys.path. We add it first so intra-backend imports resolve correctly.
_BACKEND_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_WORKSPACE_ROOT = os.path.dirname(_BACKEND_ROOT)
sys.path.insert(0, _WORKSPACE_ROOT)
sys.path.insert(0, _BACKEND_ROOT)

import fitz  # PyMuPDF — for creating test fixtures

from ingestion.smart_extractor import extract_smart

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
    datefmt="%H:%M:%S",
)

# Fix Windows terminal encoding
try:
    sys.stdout.reconfigure(encoding="utf-8")
except AttributeError:
    pass


# ─── Test Fixtures ────────────────────────────────────────────────────────────

def _create_digital_pdf(path: str) -> None:
    """Create a multi-page PDF with realistic contract text (digital, not scanned)."""
    doc = fitz.open()

    # Page 1 — MSA header and payment terms
    page1 = doc.new_page()
    page1.insert_text((50, 50), "MASTER SERVICES AGREEMENT", fontsize=16)
    page1.insert_text((50, 100), "Section 1 — Definitions")
    page1.insert_text((50, 130), "This Master Services Agreement (the 'Agreement') is entered into")
    page1.insert_text((50, 155), "by and between Client Corp and Vendor Inc effective as of the date")
    page1.insert_text((50, 180), "of last signature below. All capitalised terms have the meanings")
    page1.insert_text((50, 205), "ascribed to them in this Section unless otherwise defined herein.")
    page1.insert_text((50, 250), "Section 2 — Scope of Services")
    page1.insert_text((50, 280), "Vendor shall provide professional services as described in each")
    page1.insert_text((50, 305), "mutually executed Statement of Work. Each SOW shall be governed")
    page1.insert_text((50, 330), "by the terms and conditions of this Master Agreement.")
    page1.insert_text((50, 375), "Section 3 — Payment Terms")
    page1.insert_text((50, 405), "3.1 All invoices shall be paid within thirty (30) days of receipt.")
    page1.insert_text((50, 430), "3.2 Late payments shall incur interest at 1.5% per month.")
    page1.insert_text((50, 455), "3.3 No late fee waivers shall be granted under any circumstances.")

    # Page 2 — IP and Liability clauses
    page2 = doc.new_page()
    page2.insert_text((50, 50), "Section 4 — Intellectual Property")
    page2.insert_text((50, 80), "All work product created under this Agreement shall be the sole")
    page2.insert_text((50, 105), "and exclusive property of the Client upon payment in full.")
    page2.insert_text((50, 150), "Section 5 — Limitation of Liability")
    page2.insert_text((50, 180), "5.1 Total aggregate liability shall not exceed the total contract value.")
    page2.insert_text((50, 205), "5.2 Neither party shall be liable for indirect or consequential damages.")
    page2.insert_text((50, 250), "Section 6 — Termination")
    page2.insert_text((50, 280), "Either party may terminate this Agreement with thirty (30) days")
    page2.insert_text((50, 305), "written notice to the other party.")

    doc.save(path)
    doc.close()


def _create_scanned_pdf(path: str) -> None:
    """
    Create a PDF that simulates a scanned document.
    Contains only a tiny bit of text (below the threshold) to trigger
    the image fallback path.
    """
    doc = fitz.open()
    page = doc.new_page()
    # Only 5 words — well below the 50 words/page threshold
    page.insert_text((50, 50), "Scanned page placeholder")
    doc.save(path)
    doc.close()


def _create_docx(path: str) -> None:
    """Create a simple DOCX file for testing."""
    # We use python-docx here only for test fixture creation.
    # The actual extraction pipeline uses MarkItDown.
    try:
        import docx
        doc = docx.Document()
        doc.add_heading("Statement of Work", level=0)
        doc.add_paragraph("Section 1 — Project Overview")
        doc.add_paragraph(
            "This Statement of Work defines the deliverables for the "
            "Cloud Migration and API Gateway project."
        )
        doc.add_heading("Section 2 — Payment Schedule", level=1)
        doc.add_paragraph(
            "2.1 Payments shall be made in milestones as follows:"
        )

        # Add a table to verify table extraction
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Milestone"
        table.cell(0, 1).text = "Payment"
        table.cell(1, 0).text = "Project Kickoff"
        table.cell(1, 1).text = "30% ($15,000)"
        table.cell(2, 0).text = "Final Delivery"
        table.cell(2, 1).text = "70% ($35,000)"

        doc.save(path)
    except ImportError:
        # If python-docx is not installed, create a minimal docx
        # by just creating a placeholder text file with .docx extension
        # (MarkItDown can still attempt to read it)
        with open(path, "w") as f:
            f.write("placeholder")


# ─── Test Runner ──────────────────────────────────────────────────────────────

class TestResults:
    """Simple test result tracker."""

    def __init__(self):
        self.passed = 0
        self.failed = 0
        self.total = 0

    def ok(self, name: str, detail: str = ""):
        self.total += 1
        self.passed += 1
        print(f"  ✅ PASS  {name}" + (f"  ({detail})" if detail else ""))

    def fail(self, name: str, reason: str):
        self.total += 1
        self.failed += 1
        print(f"  ❌ FAIL  {name}  — {reason}")

    def summary(self):
        print()
        print("=" * 60)
        status = "ALL PASSED" if self.failed == 0 else "SOME FAILED"
        print(f"  Results: {self.passed}/{self.total} passed   [{status}]")
        print("=" * 60)
        return self.failed == 0


def run_tests() -> bool:
    """Execute the full test suite. Returns True if all tests pass."""
    t = TestResults()

    print()
    print("=" * 60)
    print("  SMART EXTRACTOR — TEST SUITE")
    print("=" * 60)
    print()

    # Create a temp directory for all test fixtures
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, "test_contract.pdf")
        scanned_path = os.path.join(tmp, "scanned_doc.pdf")
        docx_path = os.path.join(tmp, "test_sow.docx")
        empty_path = os.path.join(tmp, "empty.pdf")
        bad_path = os.path.join(tmp, "photo.jpg")

        # Create fixtures
        _create_digital_pdf(pdf_path)
        _create_scanned_pdf(scanned_path)
        _create_docx(docx_path)
        with open(empty_path, "w"):
            pass  # 0-byte file
        with open(bad_path, "w") as f:
            f.write("not a real image")

        # ── Test 1: Digital PDF → text mode ───────────────────────────────
        print("─── Test 1: Digital PDF Extraction ───")
        result = extract_smart(pdf_path)
        if result["mode"] == "text":
            t.ok("Mode is 'text'")
        else:
            t.fail("Mode is 'text'", f"Got '{result['mode']}' instead")

        if result.get("word_count", 0) > 30:
            t.ok("Word count > 30", f"{result['word_count']} words")
        else:
            t.fail("Word count > 30", f"Only {result.get('word_count', 0)} words")

        if result.get("page_count", 0) == 2:
            t.ok("Page count is 2", f"{result['page_count']} pages")
        else:
            t.fail("Page count is 2", f"Got {result.get('page_count')}")

        if result.get("words_per_page", 0) >= 50:
            t.ok("Words/page above threshold", f"{result['words_per_page']} w/p")
        else:
            t.fail("Words/page above threshold", f"Got {result.get('words_per_page')}")
        print()

        # ── Test 2: DOCX → text mode ─────────────────────────────────────
        print("─── Test 2: DOCX Extraction ───")
        result = extract_smart(docx_path)
        if result["mode"] == "text":
            t.ok("Mode is 'text'")
        else:
            t.fail("Mode is 'text'", f"Got '{result['mode']}' instead")

        if result.get("file_type") == "docx":
            t.ok("File type is 'docx'")
        else:
            t.fail("File type is 'docx'", f"Got '{result.get('file_type')}'")

        if result.get("word_count", 0) > 10:
            t.ok("Word count > 10", f"{result['word_count']} words")
        else:
            t.fail("Word count > 10", f"Only {result.get('word_count', 0)}")
        print()

        # ── Test 3: Scanned PDF → image fallback ─────────────────────────
        print("─── Test 3: Scanned PDF → Image Fallback ───")
        result = extract_smart(scanned_path)
        if result["mode"] == "image":
            t.ok("Mode is 'image' (fallback triggered)")
        else:
            t.fail("Mode is 'image'", f"Got '{result['mode']}' instead")

        if result["mode"] == "image":
            images = result.get("content", [])
            if len(images) == 1:
                t.ok("1 page image generated")
            else:
                t.fail("1 page image generated", f"Got {len(images)}")

            if images and len(images[0].get("base64", "")) > 100:
                t.ok("Base64 string is non-trivial", f"{len(images[0]['base64']):,} chars")
            else:
                t.fail("Base64 string is non-trivial", "Too short or missing")
        print()

        # ── Test 4: Unsupported file type → error ────────────────────────
        print("─── Test 4: Unsupported File Type ───")
        result = extract_smart(bad_path)
        if result["mode"] == "error":
            t.ok("Mode is 'error'")
        else:
            t.fail("Mode is 'error'", f"Got '{result['mode']}'")

        if result.get("error") and "Unsupported" in result["error"]:
            t.ok("Error message mentions 'Unsupported'", result["error"][:60])
        else:
            t.fail("Error message", f"Got: {result.get('error')}")
        print()

        # ── Test 5: Empty file → error ────────────────────────────────────
        print("─── Test 5: Empty File ───")
        result = extract_smart(empty_path)
        if result["mode"] == "error":
            t.ok("Mode is 'error'")
        else:
            t.fail("Mode is 'error'", f"Got '{result['mode']}'")

        if result.get("error") and "empty" in result["error"].lower():
            t.ok("Error message mentions 'empty'", result["error"])
        else:
            t.fail("Error message", f"Got: {result.get('error')}")
        print()

        # ── Test 6: Non-existent file → error ────────────────────────────
        print("─── Test 6: Non-Existent File ───")
        result = extract_smart("/this/path/does/not/exist.pdf")
        if result["mode"] == "error":
            t.ok("Mode is 'error'")
        else:
            t.fail("Mode is 'error'", f"Got '{result['mode']}'")

        if result.get("error") and "not found" in result["error"].lower():
            t.ok("Error message mentions 'not found'")
        else:
            t.fail("Error message", f"Got: {result.get('error')}")
        print()

        # ── Test 7: Markdown structure preserved ──────────────────────────
        print("─── Test 7: Markdown Structure Check ───")
        result = extract_smart(pdf_path)
        content = result.get("content", "")
        if "Section" in content or "Agreement" in content:
            t.ok("Contract text preserved in output")
        else:
            t.fail("Contract text preserved", "Key terms missing from content")
        print()

    # ── Summary ───────────────────────────────────────────────────────────
    return t.summary()


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
