"""Generate Sample_MSA.docx and Sample_SOW.docx from the .txt sources.

Run once from backend/ (or from backend/demo/):
    python demo/create_sample_docs.py

The .txt files are the source of truth; this script converts them to DOCX
so the real smart_extractor (which only accepts .docx/.pdf) can read them.
"""

import os
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")


def _txt_to_docx(txt_path: str, docx_path: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(txt_path, encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        stripped = line.rstrip("\n")

        # Skip "===" divider lines
        if stripped and all(c == "=" for c in stripped):
            continue

        # "Section X — Title" → use Word Heading 1 style so MarkItDown outputs "## Section X"
        if stripped.startswith("Section ") and "—" in stripped:
            doc.add_heading(stripped, level=1)

        # Sub-clause "4.1  ..." → Heading 2 so MarkItDown outputs "### 4.1"
        elif stripped and stripped[0].isdigit() and "." in stripped[:4] and len(stripped) > 6:
            first_token = stripped.split()[0].rstrip(".")
            if first_token.replace(".", "").isdigit():
                doc.add_heading(stripped, level=2)
            else:
                doc.add_paragraph(stripped)

        # Everything else → normal paragraph
        else:
            doc.add_paragraph(stripped)

    doc.save(docx_path)
    print(f"  Created: {docx_path}")


def main() -> None:
    pairs = [
        ("Sample_MSA.txt", "Sample_MSA.docx"),
        ("Sample_SOW.txt", "Sample_SOW.docx"),
    ]
    for txt_name, docx_name in pairs:
        txt_path = os.path.join(_HERE, txt_name)
        docx_path = os.path.join(_HERE, docx_name)
        if not os.path.isfile(txt_path):
            print(f"  SKIP (not found): {txt_path}")
            continue
        _txt_to_docx(txt_path, docx_path)
    print("Done.")


if __name__ == "__main__":
    main()
